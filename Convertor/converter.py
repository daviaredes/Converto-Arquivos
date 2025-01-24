import os  # Importa o módulo 'os' para interagir com o sistema operacional.
import csv # Importa o módulo 'csv' para trabalhar com arquivos CSV.
import fitz  # Importa o módulo 'fitz' (PyMuPDF) para trabalhar com arquivos PDF.
import subprocess  # Importa o módulo 'subprocess' para executar comandos externos.
from docx import Document  # Importa o módulo 'docx' para trabalhar com arquivos DOCX.
from reportlab.pdfgen import canvas  # Importa 'canvas' do ReportLab para criar PDFs.
from reportlab.lib.pagesizes import letter  # Importa 'letter' para o tamanho da página padrão.

# Obtém o diretório onde o script está localizado
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Definição de pastas.
PASTA_ENTRADA = os.path.join(BASE_DIR, "Arquivo a modificar")  # Pasta de entrada dos arquivos a serem convertidos.
PASTA_SAIDA = os.path.join(BASE_DIR, "Arquivos modificados")  # Pasta onde os arquivos convertidos serão salvos.

def Criacao_PDF_Texto(text, arquivo_saida):
    """Cria um arquivo PDF com o texto fornecido."""
    try:
        c = canvas.Canvas(arquivo_saida, pagesize=letter) # Cria um objeto canvas para desenhar no PDF
        textobject = c.beginText() # Inicializa um objeto de texto
        textobject.setTextOrigin(20, 750)  # Define a posição inicial do texto
        textobject.setFont("Helvetica", 14) # Define a fonte e tamanho do texto
        lines = text.splitlines() # Divide o texto em linhas
        for line in lines:
           textobject.textLine(line)  # Adiciona cada linha de texto
        c.drawText(textobject) # Desenha o texto no canvas
        c.save() # Salva o arquivo PDF
    except Exception as e:
        raise Exception(f"Erro ao criar PDF: {e}") # Lança uma exceção caso ocorra algum erro


def Leitura_Extracao(arquivo_entrada):
    """Lê um arquivo DOCX ou CSV e extrai o texto."""
    try:
        if arquivo_entrada.lower().endswith(".docx"): # Se o arquivo for .docx
            doc = Document(arquivo_entrada) # Abre o documento docx
            text = "" # Inicializa a variável de texto
            for paragraph in doc.paragraphs: # Itera sobre os paragrafos
                text += paragraph.text + '\n' # adiciona o texto do paragrafo ao texto
            return text # Retorna o texto
        elif arquivo_entrada.lower().endswith(".csv"): # Se o arquivo for csv
            with open(arquivo_entrada, 'r', encoding='utf-8') as f: # Abre o arquivo .csv
                reader = csv.reader(f) # Lê o arquivo csv
                text = "" # Inicializa a variável de texto
                for row in reader: # Itera sobre as linhas
                    text += '; '.join(row) + '\n' # adiciona a linha ao texto
                return text # Retorna o texto
        elif arquivo_entrada.lower().endswith(".pdf"):
            try:
                doc = fitz.open(arquivo_entrada)  # Abre o documento PDF.
                text = ""  # Inicializa a variável para armazenar o texto.
                for page in doc:  # Itera sobre as páginas do PDF.
                    text += page.get_text()  # Adiciona o texto de cada página.
                return text  # Retorna o texto extraído.
            except Exception as e:  # Captura qualquer erro durante a extração do texto.
                raise Exception(f"Erro ao extrair texto de PDF: {e}")  # Re-lança a exceção com mensagem de erro.
        else:
            print(f"Formato de arquivo não 8tx86suportado: {arquivo_entrada}") # Levanta uma exceção caso o arquivo não seja suportado
    except Exception as e:
        print(f"Erro ao ler arquivo: {e}") # Lança uma exceção caso ocorra algum erro

def Convertor_PDF(arquivo_entrada, arquivo_saida):
    """Converte um arquivo de texto, TXT ou PDF para PDF."""
    try:
        if arquivo_entrada.lower().endswith(('.txt')):  # Verifica se é um arquivo .txt 
            text = Leitura_Extracao(arquivo_entrada)  # Lê o texto do arquivo.
            Criacao_PDF_Texto(text, arquivo_saida)  # Cria o PDF a partir do texto.

        elif arquivo_entrada.lower().endswith('.docx'):
            text = Leitura_Extracao(arquivo_entrada)
            Criacao_PDF_Texto(text, arquivo_saida)

        elif arquivo_entrada.lower().endswith('.csv'):
            text = Leitura_Extracao(arquivo_entrada)
            Criacao_PDF_Texto(text, arquivo_saida)
    
    
    except Exception as e:  # Captura qualquer erro durante a conversão para PDF.
        print(f"Erro ao converter {arquivo_entrada}: {e}")  # Exibe mensagem de erro no console.

def Convertor_CSV(arquivo_entrada, arquivo_saida):
    """Converte um arquivo para CSV."""
    try:
        if arquivo_entrada.lower().endswith(('.txt', '.pdf', '.docx')):  # Verifica se é txt, pdf ou docx
            if arquivo_entrada.lower().endswith(('.txt')):  # Se for txt
                text = Leitura_Extracao(arquivo_entrada)  # Le o texto do arquivo
            elif arquivo_entrada.lower().endswith(('.pdf')):  # Se for pdf
               text = Leitura_Extracao(arquivo_entrada)  # Extrai o texto do pdf
            elif arquivo_entrada.lower().endswith(('.docx')): # Se for docx
                document = Document(arquivo_entrada) # Abre o documento docx
                text = "" # Inicializa o texto
                for paragraph in document.paragraphs: # Itera sobre todos os paragrafos
                   text += paragraph.text + '\n'  # Adiciona os textos dos paragrafos ao texto
            if not arquivo_entrada.lower().endswith('.csv'):
                # Cria um CSV simples com o texto
                with open(arquivo_saida, 'w+', newline='', encoding='utf-8') as f_csv:
                    f_csv.write(text)
            else:
                print(f"O arquivo de entrada '{arquivo_entrada}' já é um CSV. Nenhuma conversão necessária.")
                return  # Sai da função sem converter
        else:
            # Para outros formatos, usa o LibreOffice
            if not arquivo_entrada.lower().endswith('.csv'):
                subprocess.run(['soffice', '--headless', '--convert-to', 'csv', '--outdir', os.path.dirname(arquivo_saida), arquivo_entrada], check=True)
            else: 
                return  # Sai da função sem converter
    except Exception as e:
        print(f"Errjhvuo ao converter para CSV: {e}")

def Convertor_DOCX(arquivo_entrada, arquivo_saida):
    try:
        documento = Document()  # Cria um novo documento DOCX.

        if arquivo_entrada.lower().endswith('.txt'):
            with open(arquivo_entrada, 'r', encoding='utf-8') as f:
                text = f.read() # Lê todo o conteúdo do arquivo de texto.
            documento.add_paragraph(text) # Adiciona o texto como um parágrafo ao documento.
            documento.save(arquivo_saida) # Salva o documento como um arquivo DOCX.

        elif arquivo_entrada.lower().endswith('.pdf'):
            text = Leitura_Extracao(arquivo_entrada) # Extrai o texto do arquivo PDF.
            documento.add_paragraph(text) # Adiciona o texto como um parágrafo ao documento.
            documento.save(arquivo_saida) # Salva o documento como um arquivo DOCX.

        elif arquivo_entrada.lower().endswith('.csv'):
            with open(arquivo_entrada, 'r', encoding='utf-8') as f:
                reader = csv.reader(f) # Lê o arquivo CSV.
                for row in reader:
                    documento.add_paragraph('; '.join(row))  # Adiciona cada linha como um parágrafo, separando as colunas por ';'.
            documento.save(arquivo_saida) # Salva o documento como um arquivo DOCX.
    except Exception as e:
         print(f"Ocorreu um erro ao converter o arquivo: {e}") # Imprime uma mensagem de erro caso ocorra uma exceção.

def main():
    """Função principal para processar todos os arquivos na pasta de entrada."""
    if not os.path.exists(PASTA_ENTRADA):  # Verifica se a pasta de entrada existe.
        print(f"A pasta de entrada não existe: {PASTA_ENTRADA}")  # Exibe mensagem de erro se não existir.
        return  # Encerra a função se a pasta de entrada não existe.

    if not os.path.exists(PASTA_SAIDA):  # Verifica se a pasta de saída existe.
        os.makedirs(PASTA_SAIDA)  # Cria a pasta de saída se não existir.
        print(f"Pasta de saída criada: {PASTA_SAIDA}")  # Exibe mensagem de que a pasta de saída foi criada.

    for nome_arquivo in os.listdir(PASTA_ENTRADA):  # Itera sobre os arquivos na pasta de entrada.
        arquivo_entrada = os.path.join(PASTA_ENTRADA, nome_arquivo)  # Cria o caminho completo para o arquivo de entrada.
        if os.path.isfile(arquivo_entrada):  # Verifica se é um arquivo.
            nome_base, ext = os.path.splitext(nome_arquivo)  # Separa o nome base e a extensão do arquivo.
            arquivo_saida_pdf = os.path.join(PASTA_SAIDA, f"{nome_base}.pdf")  # Cria o caminho para o arquivo PDF de saída.
            arquivo_saida_csv = os.path.join(PASTA_SAIDA, f"{nome_base}.csv")  # Cria o caminho para o arquivo CSV de saída.
            arquivo_saida_docx = os.path.join(PASTA_SAIDA, f"{nome_base}.docx")  # Cria o caminho para o arquivo DOCX de saída.

            Convertor_PDF(arquivo_entrada, arquivo_saida_pdf)  # Converte o arquivo para .PDF.
            Convertor_CSV(arquivo_entrada, arquivo_saida_csv)  # Converte o arquivo para .CSV.
            Convertor_DOCX(arquivo_entrada, arquivo_saida_docx)  # Converte o arquivo para .DOCX.

if __name__ == "__main__":
    main()  # Chama a função principal se o script for executado diretamente.
